"""
Wordファイル処理モジュール
Word（.docx、.doc）ファイルの読み込みと処理を行います
"""
import pandas as pd
import traceback
import re
from io import BytesIO
from ..database import ensure_string

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docxがインストールされていません。pip install python-docxを実行してください。")

try:
    import olefile
    import struct
    OLEFILE_AVAILABLE = True
except ImportError:
    OLEFILE_AVAILABLE = False
    print("olefileがインストールされていません。pip install olefileを実行することを推奨します。")

def extract_text_from_docx(contents: bytes) -> str:
    """DOCXファイルからテキストを抽出する"""
    if not DOCX_AVAILABLE:
        return "[エラー: python-docxライブラリが必要です]"
    
    try:
        # BytesIOオブジェクトを作成
        docx_file = BytesIO(contents)
        
        # Documentオブジェクトを作成
        doc = Document(docx_file)
        
        extracted_text = ""
        
        # 段落のテキストを抽出
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                extracted_text += paragraph.text + "\n"
        
        # テーブルのテキストを抽出
        for table in doc.tables:
            extracted_text += "\n=== 表 ===\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    extracted_text += " | ".join(row_text) + "\n"
        
        return extracted_text.strip()
        
    except Exception as e:
        print(f"DOCXテキスト抽出エラー: {str(e)}")
        return f"[DOCXテキスト抽出エラー: {str(e)}]"

def extract_text_from_doc_basic(contents: bytes) -> str:
    """DOCファイルから基本的なテキスト抽出を試行する（簡易版）"""
    try:
        # バイナリデータから可能な限りテキストを抽出
        # これは完璧ではないが、基本的なテキストは抽出できる
        
        # 文字列として解釈可能な部分を抽出
        text_parts = []
        
        # UTF-8として読める部分を探す
        try:
            # バイナリデータを小さなチャンクに分割して処理
            chunk_size = 1024
            for i in range(0, len(contents), chunk_size):
                chunk = contents[i:i + chunk_size]
                try:
                    # UTF-8デコードを試行
                    decoded = chunk.decode('utf-8', errors='ignore')
                    # 印刷可能な文字のみを抽出
                    printable = ''.join(char for char in decoded if char.isprintable() or char.isspace())
                    if printable.strip():
                        text_parts.append(printable)
                except:
                    continue
        except Exception as e:
            print(f"基本テキスト抽出エラー: {str(e)}")
        
        # Latin-1エンコーディングでも試行
        try:
            decoded_latin1 = contents.decode('latin-1', errors='ignore')
            # ASCII文字のみを抽出
            ascii_text = ''.join(char for char in decoded_latin1 
                               if ord(char) < 128 and (char.isprintable() or char.isspace()))
            # 意味のある単語を含む行のみを抽出
            lines = ascii_text.split('\n')
            meaningful_lines = []
            for line in lines:
                clean_line = line.strip()
                if len(clean_line) > 2 and any(char.isalpha() for char in clean_line):
                    meaningful_lines.append(clean_line)
            
            if meaningful_lines:
                text_parts.extend(meaningful_lines)
                
        except Exception as e:
            print(f"Latin-1テキスト抽出エラー: {str(e)}")
        
        # 結果をまとめる
        if text_parts:
            combined_text = '\n'.join(text_parts)
            # 重複行を除去
            lines = list(dict.fromkeys(combined_text.split('\n')))
            # 短すぎる行や意味のない行を除去
            filtered_lines = [line for line in lines 
                            if len(line.strip()) > 3 and any(char.isalpha() for char in line)]
            return '\n'.join(filtered_lines)
        else:
            return "[DOCファイルからテキストを抽出できませんでした]"
            
    except Exception as e:
        print(f"DOC基本テキスト抽出エラー: {str(e)}")
        return f"[DOC基本テキスト抽出エラー: {str(e)}]"

def split_text_into_sections(text: str, filename: str) -> tuple:
    """テキストをセクションに分割する"""
    sections = {}
    
    # 見出しパターンを検出
    heading_patterns = [
        r'^(?:\d+\.?\s+)(.+)$',  # 1. タイトル
        r'^(?:[IVXLC]+\.?\s+)(.+)$',  # I. タイトル（ローマ数字）
        r'^(?:[A-Z]\.?\s+)(.+)$',  # A. タイトル
        r'^(?:第\d+[章節条項])\s*(.*)$',  # 第1章 タイトル
        r'^(?:●|■|▲|◆)\s*(.+)$',  # ● タイトル
        r'^(?:\*+|\#+)\s*(.+)$',  # * タイトル、# タイトル
    ]
    
    lines = text.split('\n')
    current_section = "概要"
    current_content = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 見出しかどうかをチェック
        is_heading = False
        for pattern in heading_patterns:
            match = re.match(pattern, line)
            if match:
                # 前のセクションを保存
                if current_content:
                    content_text = '\n'.join(current_content)
                    sections[current_section] = content_text
                
                # 新しいセクションを開始
                current_section = match.group(1).strip() if match.group(1).strip() else line
                current_content = []
                is_heading = True
                break
        
        if not is_heading:
            current_content.append(line)
    
    # 最後のセクションを保存
    if current_content:
        content_text = '\n'.join(current_content)
        sections[current_section] = content_text
    
    # セクションが作成されなかった場合は全体を一つのセクションにする
    if not sections:
        sections["全体"] = text
    
    return sections

async def process_word_file(contents: bytes, filename: str):
    """Wordファイルを処理してデータフレーム、セクション、テキストを返す"""
    try:
        print(f"Wordファイル処理開始: {filename}")
        
        # ライブラリ不足の警告
        if not DOCX_AVAILABLE:
            print("警告: python-docxライブラリが不足しています。DOCX処理精度が低下する可能性があります。")
        
        # ファイル拡張子を確認
        file_extension = filename.lower().split('.')[-1]
        
        extracted_text = ""
        
        if file_extension == 'docx':
            # DOCX形式の処理
            print("DOCX形式として処理中...")
            extracted_text = extract_text_from_docx(contents)
            
        elif file_extension == 'doc':
            # DOC形式の処理
            print("DOC形式として処理中...")
            extracted_text = extract_text_from_doc_basic(contents)
            
        else:
            extracted_text = f"[未対応のWord形式: {file_extension}]"
        
        # テキストの検証
        if not extracted_text or extracted_text.startswith('[エラー:') or extracted_text.startswith('[DOC'):
            print("通常のテキスト抽出に失敗、フォールバック処理を実行")
            # フォールバック: 基本的なバイナリ解析
            extracted_text = extract_text_from_doc_basic(contents)
        
        # セクションに分割
        sections = split_text_into_sections(extracted_text, filename)
        
        # フォーマットされたテキストを作成
        formatted_text = f"=== ファイル: {filename} ===\n\n"
        for section_name, section_content in sections.items():
            formatted_text += f"=== {section_name} ===\n{section_content}\n\n"
        
        # データフレーム用のデータを作成
        result_data = []
        for section_name, section_content in sections.items():
            if section_content.strip():
                result_data.append({
                    'section': ensure_string(section_name),
                    'content': ensure_string(section_content),
                    'source': 'Word',
                    'file': filename,
                    'url': None
                })
        
        # データがない場合のフォールバック
        if not result_data:
            result_data.append({
                'section': "文書内容",
                'content': ensure_string(extracted_text) if extracted_text else "テキストを抽出できませんでした",
                'source': 'Word',
                'file': filename,
                'url': None
            })
        
        result_df = pd.DataFrame(result_data)
        
        # すべての列の値を文字列に変換
        for col in result_df.columns:
            result_df[col] = result_df[col].apply(ensure_string)
        
        print(f"Word処理完了: {len(result_df)} セクション、{len(extracted_text)} 文字")
        return result_df, sections, formatted_text
        
    except Exception as e:
        print(f"Wordファイル処理エラー: {str(e)}")
        print(traceback.format_exc())
        
        # エラーが発生しても最低限のデータを返す
        error_message = f"Wordファイル処理中にエラーが発生しました: {str(e)}"
        empty_df = pd.DataFrame({
            'section': ["エラー"],
            'content': [error_message],
            'source': ['Word'],
            'file': [filename],
            'url': [None]
        })
        empty_sections = {"エラー": error_message}
        error_text = f"=== ファイル: {filename} ===\n\n=== エラー ===\n{error_message}\n\n"
        
        return empty_df, empty_sections, error_text

def is_word_file(filename: str) -> bool:
    """ファイルがWord形式かどうかを判定する"""
    word_extensions = {'.doc', '.docx'}
    return any(filename.lower().endswith(ext) for ext in word_extensions)

def check_word_dependencies() -> dict:
    """Word処理に必要な依存関係をチェックする"""
    return {
        'docx': DOCX_AVAILABLE,
        'olefile': OLEFILE_AVAILABLE,
        'pandas': True,  # 基本的に利用可能
    }